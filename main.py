# --------------------------------------------------------------
# Excel to word
# praveen@
# 31 jul 2022
# --------------------------------------------------------------
import pandas as pd
import docx
import configparser


def read_excel(file_name, sheet_name):
    return pd.read_excel(file_name, sheet_name)


def search(df, text):
    return df[df.apply(lambda row: row.astype(str).str.contains(text, case=False).any(), axis=1)]


def get_cols(df, cols):
    return df[cols]


def write_doc(df, file_name):

    doc = docx.Document()
    output_file_heading = config['DEFAULT']['output_file_heading']
    doc.add_heading(output_file_heading)

    headers = list(df.columns)
    h_len = len(headers)
    # print(df.values.tolist())

    for l in df.values.tolist():
        counter = 0
        table = doc.add_table(rows=h_len, cols=2, style='Table Grid')
        for h in range(h_len):
            row = table.rows[counter]
            # print(headers[h], l[h])
            row.cells[0].text = headers[h]
            row.cells[1].text = str(l[h])
            counter = counter + 1
        doc.add_paragraph() # for empty space

    try:
        doc.save(file_name)
    except PermissionError:
        print('File', file_name, 'do not have permission to write!!')


# Configuration
config = configparser.ConfigParser()
config.read('config/application.ini')

xl_file = config['DEFAULT']['xl_file']
sheet = config['DEFAULT']['xl_file_sheet']
all_df = read_excel(xl_file, sheet)

search_srt = config['DEFAULT']['search']
search_df = search(all_df, search_srt)

# print(search_df)

output_columns_str = config['DEFAULT']['output_file_columns']
output_columns = output_columns_str.split(',')
final_df = get_cols(search_df, output_columns)

output_file = config['DEFAULT']['output_file']
write_doc(final_df, output_file)
