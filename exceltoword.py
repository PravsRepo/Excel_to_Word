# --------------------------------------------------------------
# Excel to word
# praveen@
# 31 jul 2022
# --------------------------------------------------------------
import pandas as pd
import docx
import configparser
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Mm


# reading Excel sheet using pandas lib
def read_excel(file_name, sheet_name):
    try:
        return pd.read_excel(file_name, sheet_name)
    except FileNotFoundError:
        print("File not found", file_name)


# searching a string in the df and return the df
def search(df, text):
    try:
        mask = df['Name_of_college'].notnull()
        return df[mask & df['Name_of_college'].str.contains(text, case=False)]
    except Exception:
        print('Given string is not available')


# filtering the columns in the final_df
def get_cols(df, cols):
    return df[cols]


# create word document by iterating final_df
def write_doc(df, file_name):
    doc = docx.Document()
    section = doc.sections[0]
    section.header_distance = Mm(15)
    section.footer_distance = Mm(15)
    header = section.header.paragraphs[0]
    output_file_heading = config['DEFAULT']['output_file_heading']
    header.text = output_file_heading
    doc.settings.odd_and_even_pages_header_footer = True

    headers = list(df.columns)
    h_len = len(headers)

    table = doc.add_table(rows=h_len, cols=2, style='Table Grid')  # create table
    table.style = 'Light Grid Accent 2'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # print(df.values.tolist())
    # The for loop iterating the values of the df and fill them in a Word doc
    for v in df.values.tolist():
        print(df.values.tolist())
        counter = 0
        # Iterating and the rows in the table and printing the values in the table
        for h in range(h_len):
            row = table.rows[counter]
            # print(headers[h], v[h])
            row.cells[0].text = headers[h]
            row.cells[1].text = str(v[h])
            print(v[h])
            counter = counter + 1
            table.allow_autofit = True

    # doc.add_paragraph() # for empty space
    # save the document
    try:
        doc.save(file_name)
    except PermissionError:
        print('File', file_name, 'do not have permission to write!!')


# Configuration
config = configparser.ConfigParser()
config.read('config/configure.ini')

xl_file = config['DEFAULT']['xl_file']
sheet = config['DEFAULT']['xl_file_sheet']
all_df = read_excel(xl_file, sheet)

search_str = config['DEFAULT']['search']
search_df = search(all_df, search_str)

# print(search_df)
final_df = get_cols(search_df, search_df.columns[1:])

output_file = config['DEFAULT']['output_file']
write_doc(final_df, output_file)
